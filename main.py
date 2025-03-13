from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '111', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')




# import os
# import time
# import win32print
# import win32api
# from docx import Document

# # Define the document to print
# DOC_PATH = os.path.abspath("demo.docx")

# def is_printer_connected():
#     """Check if a printer is connected."""
#     printers = win32print.EnumPrinters(win32print.PRINTER_ENUM_LOCAL | win32print.PRINTER_ENUM_CONNECTIONS)
#     return len(printers) > 0

# def print_docx(doc_path):
#     """Print the Word document using the default printer."""
#     try:
#         win32api.ShellExecute(0, "print", doc_path, None, ".", 0)
#         print(f"Printing {doc_path}...")
#     except Exception as e:
#         print(f"Error while printing: {e}")

# # Wait for printer connection
# while not is_printer_connected():
#     print("Waiting for printer connection...")
#     time.sleep(5)  # Check every 5 seconds

# print("Printer connected. Printing document...")
# print_docx(DOC_PATH)



import os
import time

DOC_PATH = os.path.abspath("demo.docx")
PDF_PATH = os.path.abspath("demo.pdf")

def convert_docx_to_pdf(doc_path, pdf_path):
    """Convert .docx to .pdf using macOS built-in 'textutil'."""
    os.system(f"textutil -convert pdf {doc_path} -output {pdf_path}")

def is_printer_connected():
    """Check if a printer is available."""
    printers = os.popen("lpstat -p").read()
    return "printer" in printers

def print_pdf(pdf_path):
    """Print the PDF file using 'lp' command."""
    os.system(f"lp {pdf_path}")

# Wait for printer connection
while not is_printer_connected():
    print("Waiting for printer connection...")
    time.sleep(5)

print("Printer connected. Converting document...")
convert_docx_to_pdf(DOC_PATH, PDF_PATH)
time.sleep(5)
print("Printing document...")
print_pdf(PDF_PATH)