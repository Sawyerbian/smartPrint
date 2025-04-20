from docx import Document
from docx.shared import Inches
from docx import Document
# Load the Word document
doc = Document("demo.docx")

# Read and print all paragraphs
for para in doc.paragraphs:
    print(para.text)
doc.tables[0].cell(0,0).text = "upated value"
print(doc.tables[0].cell(0,0).text)
doc.save("demo.docx")

import subprocess
import os

libreoffice = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

def convert_docx_to_pdf(docx_path):
    """Convert DOCX to PDF using LibreOffice"""
    subprocess.run([
        libreoffice,
        "--headless",
        "--convert-to", "pdf",
        docx_path,
        "--outdir", os.path.dirname(docx_path)
    ], check=True)
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    return pdf_path

def print_pdf(pdf_path, printer=None):
    """Send the PDF to the default or specified printer using lpr"""
    cmd = ["lpr"]
    if printer:
        cmd.extend(["-P", printer])
    cmd.append(pdf_path)
    subprocess.run(cmd, check=True)

# --- Example usage ---
docx_file = "demo.docx"

try:
    pdf_file = convert_docx_to_pdf(docx_file)
    print(f"PDF created at: {pdf_file}")
    print_pdf(pdf_file)
    print("Sent to printer.")
except subprocess.CalledProcessError as e:
    print("Error occurred:", e)