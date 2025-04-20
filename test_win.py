pip install docx2pdf pywin32

import os
from docx2pdf import convert
import win32print
import win32api

from docx import Document
from docx.shared import Inches

# Load the Word document
doc = Document("demo.docx")

# Read and print all paragraphs
for para in doc.paragraphs:
    print(para.text)
doc.tables[0].cell(0,0).text = "upated value"
print(doc.tables[0].cell(0,0).text)
doc.save("demo.docx")

# --- Paths ---
docx_path = r"C:\path\to\your\document.docx"
pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

# --- Convert DOCX to PDF ---
convert(docx_path, pdf_path)
print(f"Converted to: {pdf_path}")

# --- Print PDF ---
printer_name = win32print.GetDefaultPrinter()
print(f"Printing to: {printer_name}")

# Use Acrobat Reader to silently print (Windows)
acrobat_path = r"C:\Program Files\Adobe\Acrobat DC\Acrobat\Acrobat.exe"
print_cmd = f'"{acrobat_path}" /t "{pdf_path}" "{printer_name}"'
win32api.ShellExecute(0, "open", acrobat_path, f'/t "{pdf_path}" "{printer_name}"', ".", 0)


import subprocess

# Convert
subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "your.docx"])

# Print
subprocess.run(["lpr", "your.pdf"])